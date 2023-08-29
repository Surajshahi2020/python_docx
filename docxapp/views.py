from docx import Document as PythonDocxDocument
from django.http import HttpResponse
from docx.shared import Cm, RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


def generate_docx_table(request):
    data = [
        {
            "title": "Python Questions",
            "questions": [
                {
                    "q_title": "Machine Learning",
                    "question": "What is Machine Learning?",
                    "answer": "Machine learning is the branch of AI and computer science.",
                    "level": "Hard",
                },
                {
                    "q_title": "Machine Learning",
                    "question": "Advantages of Machine Learning?",
                    "answer": "Used to analyze large dataset for useful information.",
                    "level": "Hard",
                },
                {
                    "q_title": "Machine Learning",
                    "question": "Advantages of Machine Learning?",
                    "answer": "Used to analyze large dataset for useful information.",
                    "level": "Hard",
                },
                {
                    "q_title": "Machine Learning",
                    "question": "Advantages of Machine Learning?",
                    "answer": "Used to analyze large dataset for useful information.",
                    "level": "Hard",
                },
            ],
        },
    ]

    document = PythonDocxDocument()
    for item in data:
        title = item["title"]
        document.add_heading(title, level=1)

        document.add_paragraph()

        table = document.add_table(rows=1, cols=5)
        table.style = "Table Grid"
        header_row = table.rows[0]
        header_row.cells[0].text = "SN"
        header_row.cells[1].text = "Title"
        header_row.cells[2].text = "Question"
        header_row.cells[3].text = "Answer"
        header_row.cells[4].text = "Level"

        sn_column = table.columns[0]
        sn_column.width = Cm(1)

        for cell in header_row.cells:
            tcPr = cell._tc.get_or_add_tcPr()
            color = RGBColor(173, 216, 230)
            shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" />')
            tcPr.append(shading_elm)

        # Fill the table with question and answer data
        for idx, question in enumerate(item["questions"], start=1):
            q_title = question["q_title"]
            q_question = question["question"]
            q_answer = question["answer"]
            q_level = question["level"]

            row = table.add_row().cells
            row[0].text = str(idx)
            row[1].text = q_title
            row[2].text = q_question
            row[3].text = q_answer
            row[4].text = q_level

    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = 'attachment; filename="interview.docx"'
    document.save(response)

    return response


def generate_docx_step(request):
    data = [
        {
            "title": "Python",
            "questions": [
                {
                    "question": "What is Machine Learning?",
                    "answer": "Machine learning is the branch of AI and computer science.",
                },
                {
                    "q_title": "Advantages of Machine Learning?",
                    "question": "Advantages of Machine Learning?",
                    "answer": "Used to analyze large dataset for useful information.",
                },
                {
                    "question": "What is Machine Learning?",
                    "answer": "Machine learning is the branch of AI and computer science.",
                },
                {
                    "q_title": "Advantages of Machine Learning?",
                    "question": "Advantages of Machine Learning?",
                    "answer": "Used to analyze large dataset for useful information.",
                },
            ],
        },
    ]

    document = PythonDocxDocument()
    for item in data:
        title = item["title"]
        document.add_heading(title, level=1)

        document.add_paragraph()

        for idx, question in enumerate(item["questions"], start=1):
            q_question = question["question"]
            q_answer = question["answer"]

            document.add_paragraph(f"Question: {idx}. {q_question}", style="Heading2")
            document.add_paragraph(f"Answer: {q_answer}", style="BodyText")

            document.add_paragraph()

    response = HttpResponse(
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    response["Content-Disposition"] = 'attachment; filename="interview.docx"'
    document.save(response)

    return response
